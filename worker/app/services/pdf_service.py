from __future__ import annotations
import json
import os
import structlog

import pdfplumber
import pytesseract

from app.services.sensitivity import SensitivityScanner

logger = structlog.get_logger()

PDF_SYSTEM_PROMPT = """Kamu menganalisis dokumen PDF/Word untuk grup KKN / komunitas desa di Indonesia.

PERAN: editor + perangkum cerdas — BUKAN mesin copy-paste.
Tulis ulang isi dokumen dengan bahasa yang rapi, baku-santai, mudah dibaca di HP.
Target: paham & actionable dalam ~1 menit; BUKAN salinan mentah / notulen versi 2 yang redundant.

YANG BOLEH (disarankan):
1. Memperbaiki typo / ejaan yang jelas (contoh: INFRAKSTRUKTUR→Infrastruktur, "nyusul"→"menyusul" di kalimat formal-santai).
2. Merapikan kalimat acak-acakan jadi poin yang mudah dipindai.
3. Mengelompokkan ulang item yang campur-aduk (mis. pecah "peralatan masak" vs "kebersihan" jika di dokumen berantakan).
4. Menyeragamkan gaya: judul section singkat, bullet padat, nama orang konsisten.
5. Menyederhanakan frasa bertele-tele tanpa menghilangkan makna.
6. Menandai ketidakjelasan di dokumen (bukan mengarang jawaban), mis. "Kompor & gas: masih perlu dikonfirmasi".

YANG TIDAK BOLEH:
1. Mengarang fakta, nama, tanggal, jam, lokasi, atau penugasan yang tidak ada di dokumen.
2. Mengubah makna kesepakatan (hanya rapikan bahasa).
3. Menghapus detail penting (siapa bawa apa, kapan survei, daftar belanja inti).
4. Menyalin mentah typo/notasi jelek "asal sama" — utamakan keterbacaan + akurasi.
5. Meledakkan penugasan jadi puluhan baris "Membawa X — Nama".
6. Mengulang isi yang sama di banyak field (anti-redundansi).

STRUKTUR OUTPUT:
1. Jika dokumen ber-section (A/B/C/heading): isi `sections` sebagai INTI.
   - heading: rapikan (boleh "B. Revisi program kerja" bukan salin huruf besar acak).
   - points: tulis ulang poin dengan bahasa jernih; daftar panjang digabung per baris/kategori.
2. purpose: 2–3 kalimat gaya sendiri — apa dokumen ini & untuk apa dibaca.
3. Field lain (key_points, decisions, tasks, assignments, schedule, open_questions,
   shopping_list, deadlines): biarkan [] jika sudah tercakup di sections.
   Hanya isi jika ada "masih terbuka" yang perlu disorot terpisah.

OUTPUT JSON:
{
  "title": "judul rapi",
  "purpose": "2-3 kalimat bahasa sendiri",
  "sections": [{"heading": "A. ...", "points": ["poin ditulis ulang, jelas"]}],
  "key_points": [],
  "decisions": [],
  "tasks": [],
  "assignments": [],
  "schedule": [],
  "open_questions": [],
  "shopping_list": [],
  "deadlines": [],
  "source_pages": []
}

Balikkan HANYA JSON valid, tanpa markdown."""


class PdfService:
    def __init__(self):
        self.sensitivity_scanner = SensitivityScanner()

    async def analyze(self, request) -> dict:
        """Extract text from PDF/Word, redact sensitive data, analyze thoroughly."""
        file_path = request.file_path

        if not os.path.exists(file_path):
            return {"error": "File not found", "status": "error"}

        try:
            extracted_text, page_count = self._extract_text(file_path)

            if not extracted_text.strip():
                return {
                    "status": "unprocessable",
                    "error": "No text could be extracted from document",
                }

            safe_text, sensitivity_result = self.sensitivity_scanner.redact(extracted_text)

            analysis = await self._analyze_with_ai(
                safe_text,
                {
                    **(request.metadata or {}),
                    "page_count": page_count or (request.metadata or {}).get("page_count"),
                },
            )
            if isinstance(analysis, dict) and analysis.get("error"):
                return {"status": "error", "error": analysis["error"]}

            if isinstance(analysis, dict):
                analysis["redacted"] = sensitivity_result.is_sensitive
                if sensitivity_result.is_sensitive:
                    analysis["sensitivity_note"] = sensitivity_result.summary

            return {
                "status": "analyzed",
                "page_count": page_count,
                "analysis": analysis,
                "redacted": sensitivity_result.is_sensitive,
            }

        except Exception as e:
            logger.error("Document processing failed", error=str(e), document_id=request.document_id)
            return {"status": "error", "error": str(e)}

    def _extract_text(self, file_path: str) -> tuple[str, int]:
        lower = file_path.lower()
        if lower.endswith(".docx") or lower.endswith(".doc"):
            return self._extract_word(file_path)
        return self._extract_pdf(file_path)

    def _extract_word(self, file_path: str) -> tuple[str, int]:
        try:
            from docx import Document  # type: ignore

            doc = Document(file_path)
            parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts)
            return text, max(1, len(parts) // 40 or 1)
        except ImportError:
            logger.error("python-docx not installed")
            raise RuntimeError("Word support requires python-docx")
        except Exception as e:
            logger.error("Word extraction failed", error=str(e))
            raise

    def _extract_pdf(self, file_path: str) -> tuple[str, int]:
        text_parts = []
        page_count = 0

        try:
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Halaman {i+1}]\n{page_text}")
                    else:
                        try:
                            ocr_text = self._ocr_page(page)
                            if ocr_text:
                                text_parts.append(f"[Halaman {i+1} (OCR)]\n{ocr_text}")
                        except Exception as e:
                            logger.warning(f"OCR failed for page {i+1}", error=str(e))
        except Exception as e:
            logger.error("PDF extraction failed", error=str(e))
            raise

        return "\n\n".join(text_parts), page_count

    def _ocr_page(self, page) -> str:
        try:
            image = page.to_image(resolution=300)
            pil_image = image.original
            text = pytesseract.image_to_string(pil_image, lang="ind+eng")
            return text.strip()
        except Exception as e:
            logger.warning("OCR failed", error=str(e))
            return ""

    async def _analyze_with_ai(self, text: str, metadata: dict) -> dict:
        from app.providers.cascade import ProviderCascade

        # Keep nearly full short/medium docs; only hard-cap huge files
        max_chars = 80_000
        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n[Dokumen dipotong karena terlalu panjang — prioritaskan bagian awal & daftar penugasan]"

        user_content = (
            f"Dokumen: {metadata.get('filename', 'unknown')}\n"
            f"Jumlah halaman: {metadata.get('page_count') or '?'}\n"
            f"Panjang teks: {len(text)} karakter\n\n"
            "Tugas: baca teks di bawah, lalu TULIS ULANG ringkasan terstruktur. "
            "Rapikan bahasa & typo yang jelas, perjelas pengelompokan, "
            "tetap setia pada fakta dokumen (nama, tanggal, penugasan). "
            "Jangan salin mentah jika kalimat di sumber berantakan. "
            "Jangan buang detail penting; jangan mengarang.\n\n"
            f"Teks sumber (mungkin ada typo):\n{text}"
        )

        cascade = ProviderCascade()
        for route in cascade.get_routes("pdf"):
            try:
                result = await cascade.call(
                    route=route,
                    system_prompt=PDF_SYSTEM_PROMPT,
                    user_content=user_content,
                    response_format={"type": "json_object"},
                )
                return json.loads(result.content)
            except Exception as e:
                logger.warning("PDF route failed", route=route.id, error=str(e))
                continue

        return {"error": "All provider routes failed for PDF analysis"}
